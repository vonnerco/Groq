from __future__ import annotations

import hashlib
import json
import os
from pathlib import Path
from tempfile import TemporaryDirectory

import pandas as pd
from docx import Document
from PIL import Image, ImageDraw


def file_signature(name: str, data: bytes) -> str:
    return f"{name}|{hashlib.sha256(data).hexdigest()}"


def preview_csv(path: Path):
    try:
        return pd.read_csv(path)
    except Exception:
        return pd.read_csv(path, encoding="latin-1")


def preview_xlsx(path: Path):
    return pd.read_excel(path)


def preview_docx(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def preview_pdf(path: Path) -> str | None:
    try:
        from pypdf import PdfReader
    except Exception:
        return None
    try:
        reader = PdfReader(str(path))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception:
                return None
        chunks = []
        for page in reader.pages[:5]:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(chunks).strip()
        return text or None
    except Exception:
        return None


def main() -> int:
    with TemporaryDirectory() as td:
        root = Path(td)
        csv_path = root / "sample.csv"
        xlsx_path = root / "sample.xlsx"
        docx_path = root / "sample.docx"
        png_path = root / "sample.png"

        df = pd.DataFrame({"name": ["alpha", "beta"], "value": [1, 2]})
        df.to_csv(csv_path, index=False)
        df.to_excel(xlsx_path, index=False)

        doc = Document()
        doc.add_paragraph("Hello upload preview")
        doc.add_paragraph("Second line")
        doc.save(docx_path)

        img = Image.new("RGB", (64, 64), "white")
        draw = ImageDraw.Draw(img)
        draw.rectangle((8, 8, 56, 56), outline="black", width=2)
        img.save(png_path)

        checks = {
            "csv_rows": len(preview_csv(csv_path)),
            "xlsx_rows": len(preview_xlsx(xlsx_path)),
            "docx_text": preview_docx(docx_path),
            "png_size": Image.open(png_path).size,
            "csv_signature": file_signature(csv_path.name, csv_path.read_bytes()),
        }

        pdf_candidate = root / "sample.pdf"
        pdf_text = preview_pdf(pdf_candidate) if pdf_candidate.exists() else None
        checks["pdf_preview"] = pdf_text

        print(json.dumps(checks, indent=2, default=str))
        assert checks["csv_rows"] == 2
        assert checks["xlsx_rows"] == 2
        assert "Hello upload preview" in checks["docx_text"]
        assert checks["png_size"] == (64, 64)
        assert checks["csv_signature"].startswith("sample.csv|")
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
